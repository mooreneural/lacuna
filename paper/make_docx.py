#!/usr/bin/env python
"""Render manuscript.md to a Word document.

Kept in the repository rather than run once by hand so the .docx can be
regenerated whenever the numbers change, the same way the figures are. Pandoc
would be the obvious tool and is not available here, so this handles the subset
of Markdown the manuscript actually uses: YAML front matter, ATX headings, pipe
tables with alignment, and inline bold, italic, code and links.

  python paper/make_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

#: Text column width for US Letter with one-inch margins.
FIG_WIDTH_IN = 6.5

_IMAGE = re.compile(r"^!\[(?P<alt>[^]]*)\]\((?P<path>[^)]+)\)\s*$")

HERE = Path(__file__).resolve().parent
SRC = HERE / "manuscript.md"
OUT = HERE / "manuscript.docx"
OUT_JOURNAL = HERE / "manuscript_twocolumn.docx"

#: Two-column geometry measured from Current Opinion in Structural Biology
#: (Jing et al. 2026): 0.78 in margins, ~3.4 in columns, ~0.22 in gutter.
JOURNAL = {"margin": 0.78, "gutter": 0.22, "body_pt": 9.5, "font": "Georgia"}

#: Body font for the single-column variant. bioRxiv names Times, Times New Roman,
#: Courier, Helvetica, Arial and Symbol as the fonts it can convert reliably, and
#: warns that anything else "may not convert at all or may appear blocky". Times
#: New Roman is on that list and is also the closer match to the Computer Modern
#: that preprints in this area are typically set in, so there is nothing to trade
#: off here.
BODY_FONT = "Times New Roman"

#: Same reasoning for inline code: Courier New is on the supported list.
CODE_FONT = "Courier New"

#: Sections left unnumbered, matching the convention in arXiv-style reports where
#: the abstract and the back matter sit outside the numbered sequence.
UNNUMBERED = {"abstract", "data and code availability", "competing interests",
              "references", "acknowledgements", "funding",
              "author contributions"}


def _set_columns(section, num: int, gutter_in: float = 0.22) -> None:
    """Set the column count on a section.

    python-docx has no column API, so this edits sectPr directly. Equal-width
    columns are implied when w:equalWidth is set and no individual w:col entries
    are present, which is what Word writes for a plain two-column layout.
    """
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    for child in list(cols):
        cols.remove(child)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:equalWidth"), "1")
    cols.set(qn("w:space"), str(int(gutter_in * 1440)))


def _span(doc, columns: int):
    """Break to a new continuous section with the given column count.

    Journals set body text in two columns and let wide tables and figures span
    both. In Word that is a continuous section break either side of the wide
    element, which is what this produces.
    """
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    for attr, val in (("left_margin", JOURNAL["margin"]),
                      ("right_margin", JOURNAL["margin"]),
                      ("top_margin", 0.9), ("bottom_margin", 0.9)):
        setattr(section, attr, Inches(val))
    _set_columns(section, columns, JOURNAL["gutter"])
    return section

#: Bold, italic, inline code, and links, in one pass so nesting cannot double-apply.
_INLINE = re.compile(
    r"\*\*(?P<bold>[^*]+)\*\*"
    r"|(?<!\*)\*(?P<ital>[^*]+)\*(?!\*)"
    r"|`(?P<code>[^`]+)`"
    r"|\[(?P<text>[^]]+)\]\((?P<url>[^)]+)\)"
)


def _front_matter(lines: list[str]) -> tuple[dict, int]:
    """Parse the leading YAML block. Returns (fields, index after the block).

    Deliberately not a YAML parser: the block is three scalars plus one nested
    author record, and a dependency for that would be worse than ten lines.
    """
    if not lines or lines[0].strip() != "---":
        return {}, 0
    meta: dict[str, str] = {}
    i = 1
    while i < len(lines) and lines[i].strip() != "---":
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("- name:"):
            meta["author"] = stripped.split(":", 1)[1].strip()
        elif stripped.startswith("orcid:"):
            meta["orcid"] = stripped.split(":", 1)[1].strip()
        elif stripped.startswith("affiliation:"):
            meta["affiliation"] = stripped.split(":", 1)[1].strip()
        elif not line.startswith(" ") and ":" in stripped:
            key, val = stripped.split(":", 1)
            meta[key.strip()] = val.strip().strip('"')
        i += 1
    return meta, i + 1


def _add_runs(paragraph, text: str) -> None:
    """Append text to a paragraph, honouring inline markup."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("bold"):
            paragraph.add_run(m.group("bold")).bold = True
        elif m.group("ital"):
            paragraph.add_run(m.group("ital")).italic = True
        elif m.group("code"):
            run = paragraph.add_run(m.group("code"))
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
        else:
            # Links render as styled text: a real hyperlink relationship needs
            # raw XML in python-docx, and the URLs here are all in the
            # availability section where the text is the address anyway.
            run = paragraph.add_run(m.group("text"))
            run.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _alignments(sep: str) -> list[str]:
    out = []
    for cell in [c.strip() for c in sep.strip().strip("|").split("|")]:
        if cell.endswith(":") and cell.startswith(":"):
            out.append("center")
        elif cell.endswith(":"):
            out.append("right")
        else:
            out.append("left")
    return out


_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT,
          "right": WD_ALIGN_PARAGRAPH.RIGHT,
          "center": WD_ALIGN_PARAGRAPH.CENTER}


def _add_table(doc: Document, rows: list[str]) -> None:
    header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    align = _alignments(rows[1])
    body = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows[2:]]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text, a in zip(table.rows[0].cells, header, align):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = _ALIGN[a]
        _add_runs(p, text)
        for run in p.runs:
            run.bold = True
    for line in body:
        cells = table.add_row().cells
        for cell, text, a in zip(cells, line, align):
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = _ALIGN[a]
            _add_runs(p, text)
    doc.add_paragraph()


def _left_rule(p, color: str = "8a8983", sz: int = 12) -> None:
    """A thin rule down the left edge, the usual typographic mark for a callout."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_callout(doc, text: str) -> None:
    """A Markdown blockquote.

    Rendered as an indented paragraph with a left rule. Without this the '>'
    markers reached the document as literal text, including the ones opening
    each wrapped line, because the fallback path joins a paragraph's lines with
    spaces and has no reason to treat them as syntax.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.28)
    pf.right_indent = Inches(0.18)
    pf.space_before = Pt(7)
    pf.space_after = Pt(7)
    _add_runs(p, text)
    _left_rule(p)


def _add_code(doc, code_lines: list[str], journal: bool) -> None:
    """A fenced code block: monospace, one line per source line.

    Line breaks are explicit rather than separate paragraphs so the block stays
    together on the page, and the info string after the opening fence is dropped.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.28)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    for n, line in enumerate(code_lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5 if journal else 9.5)
        # python-docx sets only the Latin font; without this Word may substitute
        # for the other script ranges and break the monospace alignment.
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), "Consolas")
        if n < len(code_lines) - 1:
            run.add_break()


def _title_block(doc, meta, journal: bool) -> None:
    """Title, byline and affiliation.

    The affiliation was originally omitted here, on the reasoning that bioRxiv
    collects it on the submission form and reprints it on the posted page. That
    holds for the posted preprint but not for the .docx read on its own, or for
    a journal submission, so it is rendered. ORCID and date stay out: bioRxiv
    does print those, and they are kept in the YAML as the record of record.

    Both are driven from the front matter, so the rendered document cannot
    disagree with the source about who wrote it or where.
    """
    if meta.get("title"):
        h = doc.add_heading(meta["title"], level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if journal:
            for run in h.runs:
                run.font.size = Pt(17)
    if meta.get("author"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(meta["author"]).italic = True
    if meta.get("affiliation"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(meta["affiliation"])
        run.font.size = Pt(JOURNAL["body_pt"] if journal else 10)


def build(journal: bool) -> Path:
    """Render the manuscript. ``journal`` selects the two-column variant."""
    lines = SRC.read_text(encoding="utf-8").splitlines()
    meta, start = _front_matter(lines)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = JOURNAL["font"] if journal else BODY_FONT
    style.font.size = Pt(JOURNAL["body_pt"] if journal else 11)

    margin = JOURNAL["margin"] if journal else 1.0
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for attr in ("left_margin", "right_margin"):
            setattr(section, attr, Inches(margin))
        for attr in ("top_margin", "bottom_margin"):
            setattr(section, attr, Inches(0.9 if journal else 1.0))

    full_w = 8.5 - 2 * margin
    fig_w = full_w if journal else FIG_WIDTH_IN

    # Title block spans the page even in the two-column variant, as journals set it.
    _title_block(doc, meta, journal)
    if journal:
        _span(doc, 2)

    i, n_tables, n_images, n_missing = start, 0, 0, 0
    sec_no, sub_no = [0], [0]
    in_two_col = journal

    def to_full():
        """Break out to a full-width section for a wide element."""
        nonlocal in_two_col
        if journal and in_two_col:
            _span(doc, 1)
            in_two_col = False

    def to_body():
        nonlocal in_two_col
        if journal and not in_two_col:
            _span(doc, 2)
            in_two_col = True

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("#"):
            to_body()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            if level == 2:
                if text.lower() not in UNNUMBERED:
                    sec_no[0] += 1
                    sub_no[0] = 0
                    text = f"{sec_no[0]}  {text}"
            elif level == 3 and sec_no[0]:
                sub_no[0] += 1
                text = f"{sec_no[0]}.{sub_no[0]}  {text}"
            h = doc.add_heading(text, level=min(level, 4))
            if journal:
                for run in h.runs:
                    run.font.size = Pt(11.5 if level <= 2 else 10)
            i += 1
            continue

        img = _IMAGE.match(stripped)
        if img:
            to_full()
            src = (HERE / img.group("path")).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if src.exists():
                p.add_run().add_picture(str(src), width=Inches(fig_w))
                n_images += 1
            else:
                run = p.add_run(f"[ {img.group('alt')}: {img.group('path')} "
                                f"not generated yet ]")
                run.italic = True
                run.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
                n_missing += 1
            i += 1
            continue

        if (stripped.startswith("|") and i + 1 < len(lines)
                and set(lines[i + 1].strip()) <= set("|-: ")
                and "-" in lines[i + 1]):
            to_full()
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            _add_table(doc, block)
            n_tables += 1
            continue

        if stripped.startswith("```"):
            to_body()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip())
                i += 1
            i += 1                      # closing fence
            _add_code(doc, code, journal)
            continue

        if stripped.startswith(">"):
            to_body()
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            _add_callout(doc, " ".join(quote).strip())
            continue

        to_body()
        buf = []
        while (i < len(lines) and lines[i].strip()
               and not lines[i].strip().startswith(("#", "|", "![", "```", ">"))):
            buf.append(lines[i].strip())
            i += 1
        _add_runs(doc.add_paragraph(), " ".join(buf))

    out = OUT_JOURNAL if journal else OUT
    doc.save(out)
    msg = f"wrote {out}  ({n_tables} tables, {n_images} figures"
    if n_missing:
        msg += f", {n_missing} MISSING"
    print(msg + ")")
    return out


def main() -> None:
    if not SRC.exists():
        sys.exit(f"missing {SRC}")
    build(journal=False)
    build(journal=True)


if __name__ == "__main__":
    main()
